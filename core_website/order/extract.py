from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from docx import Document
from django.utils.timezone import localtime
from .models import Order


def order_word(request, order_id):
    order = get_object_or_404(Order, id=order_id)
    document = Document()

    document.add_heading(f'Order #{order.id} - {order.get_status_display()}', level=1)

    document.add_paragraph('Seller Info')

    seller_table = document.add_table(rows=6, cols=2)
    seller_table.style = 'Table Grid'
    info = [
        ('Seller', 'Core Official Store'),
        ('Seller id', '0123456789'),
        ('Website', 'core.com'),
        ('Email', 'support@coretechstore.com'),
        ('Address', '123 Core Street, Odesa, Ukraine'),
        ('Phone', '+38-067-123-45-67'),

    ]
    for i, row in enumerate(info):
        seller_table.cell(i, 0).text = row[0]
        seller_table.cell(i, 1).text = row[1]

    document.add_paragraph('\nCustomer info')

    customer_table = document.add_table(rows=7, cols=2)
    customer_table.style = 'Table Grid'
    info = [
        ('User', str(order.user)),
        ('Name', f'{order.first_name} {order.last_name}'),
        ('Total cost', f'{order.get_total_cost()/100:.2f} $'),
        ('Email', order.email),
        ('Address', f'{order.city} {order.address}'),
        ('Created', localtime(order.created).strftime("%B %d, %Y, %I:%M %p")),
        ('Paid', '✔️' if order.paid else '❌'),
    ]
    for i, row in enumerate(info):
        customer_table.cell(i, 0).text = row[0]
        customer_table.cell(i, 1).text = row[1]

    document.add_paragraph('\nProducts')
    item_table = document.add_table(rows=1, cols=5)
    item_table.style = 'Table Grid'
    hdr_cells = item_table.rows[0].cells
    hdr_cells[0].text = '#'
    hdr_cells[1].text = 'Product name'
    hdr_cells[2].text = 'Price'
    hdr_cells[3].text = 'Quantity'
    hdr_cells[4].text = 'Total cost'

    for idx, item in enumerate(order.items.all(), start=1):
        row_cells = item_table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = item.product.name
        row_cells[2].text = f'{item.price/100:.2f} $'
        row_cells[3].text = str(item.quantity)
        row_cells[4].text = f'{item.get_cost()/100:.2f} $'

    row_cells = item_table.add_row().cells
    row_cells[0].text = ''
    row_cells[1].text = 'Total order cost'
    row_cells[2].text = ''
    row_cells[3].text = ''
    row_cells[4].text = f'{order.get_total_cost()/100:.2f} $'

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename=order_{order.id}.docx'
    document.save(response)
    return response
